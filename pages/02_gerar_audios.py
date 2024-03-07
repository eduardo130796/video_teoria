import streamlit as st
from docx import Document
import io
import os
import zipfile
import elevenlabs
from elevenlabs import generate, Voice, VoiceSettings
import docx
import shutil
from elevenlabs import RateLimitError

# Função para dividir texto
# Função para dividir texto
def dividir_texto(texto, limite, separador):
    palavras = texto.split(separador)
    partes = []
    parte_atual = ""
    for palavra in palavras:
        if len(parte_atual + palavra) > limite:
            partes.append(parte_atual)
            parte_atual = ""
        parte_atual += palavra + separador
    partes.append(parte_atual)
    return partes


def streamlit_app():
    st.title("Gerador de Narração de Slides")

    # Campos para chave da API e voice_id
    api_key = st.text_input("Chave da API ElevenLabs", type="password")
    voice_id = st.text_input("Voice ID")

    # Upload do documento Word
    doc_file = st.file_uploader("Escolha o arquivo DOCX", type="docx")
    
    doc = docx.Document(doc_file)
    texto_arquivo = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    st.text_area("Texto do Arquivo", texto_arquivo, height=300)
    
    if doc_file is not None:
        # Ler o arquivo DOCX da memória
        doc_memory = io.BytesIO(doc_file.getvalue())
        doc = Document(doc_memory)

    
    # Define a chave da API do ElevenLabs
    elevenlabs.set_api_key(api_key)

    
    

    # Diretório temporário para os arquivos de áudio
    audio_dir = "audio_narracao"
    if not os.path.exists(audio_dir):
        os.mkdir(audio_dir)

# Seção para visualização dos arquivos criados
    with st.expander("Visualizar arquivos criados"):
        for root, dirs, files in os.walk(audio_dir):
            for file in files:
                audio_path = os.path.join(root, file)
                with open(audio_path, "rb") as audio_file:
                    st.audio(audio_file.read(), format="audio/mp3")
                if st.button(f"Excluir {file}"):
                    os.remove(audio_path)
                    st.success(f"Arquivo {file} excluído.")
        # Adicionar botão para limpar a pasta
        limpar_pasta_button = st.button("Limpar Pasta de Áudios")
        if limpar_pasta_button:
            try:
                shutil.rmtree(audio_dir)
                os.makedirs(audio_dir)
                st.success("Pasta de áudios limpa com sucesso.")
            except Exception as e:
                st.warning(f"Erro ao limpar pasta de áudios: {e}")
    # Botão para iniciar a criação dos áudios
    start_button = st.button("Iniciar Criação dos Áudios")
    # Botão para interromper a criação dos áudios
    stop_button = st.button("Parar Criação dos Áudios")

    
    if start_button and api_key and voice_id:

        try:
            if doc_file is not None and api_key and voice_id:
                # Processa o documento
                current_slide = 0
                file_count = 0
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text.startswith("SLIDE:"):
                        current_slide = int(text.split(': ')[1])
                        file_count = 0
                    elif text and text != '.':
                        file_count += 1
                        partes = dividir_texto(text, 200, '.')
                        for parte in partes:
                            if parte.strip() and parte.strip() != '.':
                                arquivo_audio = f"{audio_dir}/{current_slide}.{file_count}_narracao_slide.mp3"
                                if os.path.exists(arquivo_audio):
                                    st.audio(arquivo_audio, format="audio/mp3")
                                    st.write(f"O arquivo {arquivo_audio} já existe na pasta.")
                                else:
                                    audio = generate(
                                        text=parte,
                                        voice=Voice(
                                            voice_id=voice_id,
                                            settings=VoiceSettings(stability=1.0, similarity_boost=0.70, style=0.0, use_speaker_boost=True)
                                        ),
                                        model="eleven_multilingual_v2"
                                    )
                                    with open(arquivo_audio, 'wb') as file:
                                        file.write(audio)
                                    st.audio(arquivo_audio, format="audio/mp3")
                                    st.success(f"Texto criado para Slide {current_slide}, Parte {file_count}:\n\n{parte}")
                                if not parte:
                                    st.warning(f"Não foi possível processar o texto: {parte}. Não há conteúdo suficiente.")

            

                # Criar arquivo ZIP com os arquivos de áudio
                zip_memory = io.BytesIO()
                with zipfile.ZipFile(zip_memory, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for root, dirs, files in os.walk(audio_dir):
                        for file in files:
                            zipf.write(os.path.join(root, file), file)
                
                zip_memory.seek(0)

                # Botão de download do ZIP
                st.download_button(
                    label="Baixar narrações em ZIP",
                    data=zip_memory,
                    file_name="narrações.zip",
                    mime="application/zip"
                )
        except RateLimitError as e:
            st.error("Os créditos da API acabaram. É necessário trocar de conta ou aguardar a renovação dos créditos para continuar.")
   
        # Limpar o diretório de áudio
        for root, dirs, files in os.walk(audio_dir):
            for file in files:
                os.remove(os.path.join(root, file))
        os.rmdir(audio_dir)

        # Condição para interromper a criação dos áudios
    if stop_button:
        st.warning("A criação dos áudios foi interrompida.")
        return  # Interrompe a execução do script

if __name__ == "__main__":
    streamlit_app()